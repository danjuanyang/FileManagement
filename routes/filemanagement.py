# filemanagement.py
import io
import re
import shutil
import sys
import tempfile
import time
from pydoc import html

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.oxml import qn
from flask import Blueprint, Response, stream_with_context
from flask_cors import CORS
from datetime import datetime
import jwt
from pptx import Presentation
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from sqlalchemy import or_, text, func
from sqlalchemy.orm import joinedload
from werkzeug.utils import secure_filename
from flask import Blueprint, request, jsonify, send_file, abort, current_app
from flask_cors import CORS
from flask import jsonify, request
import os
from models import db, Project, ProjectFile, ProjectStage, User, StageTask, FileContent, UserActivityLog, Subproject
from auth import get_employee_id
from utils.activity_tracking import track_activity
from docx import Document

from reportlab.lib.pagesizes import A4, landscape
# 搜索

from .file_indexer import update_file_index, get_mime_type, create_file_index

from werkzeug.utils import secure_filename
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
# from docx2pdf import convert
# from win32com import client
# import pythoncom
# from PyPDF2 import PdfMerger

import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from openpyxl import load_workbook
from pptx import Presentation
from PyPDF2 import PdfMerger
from flask import send_file, Response, stream_with_context, jsonify
import time
import shutil

# 测试开发 使用本路径
# 获取Python解释器所在目录
python_dir = os.path.dirname(sys.executable)
# 设置上传目录为程序根目录下文件夹
UPLOAD_FOLDER = os.path.join(python_dir, 'uploads')
ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf', 'xls', 'xlsx', 'txt', 'zip', 'rar'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
BASE_UPLOAD_FOLDER = os.path.join(python_dir, 'uploads')  # 基础上传目录

# 群晖路径
# python_dir  = '/volume1/web/FileManagementFolder/db'
# UPLOAD_FOLDER = '/volume1/web/FileManagementFolder/uploads'
# ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf', 'xls', 'xlsx', 'txt', 'zip', 'rar'}
# MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
# BASE_UPLOAD_FOLDER = os.path.join(python_dir, 'uploads')  # 基础上传目录


files_bp = Blueprint('files', __name__)
CORS(files_bp)


def sanitize_filename(filename):
    """
    安全地处理文件名，保留中文字符
    移除或替换不安全的字符，但保留中文和基本标点
    """
    # 替换Windows文件系统中的非法字符
    illegal_chars = r'[<>:"/\\|?*\x00-\x1f]'
    # 将非法字符替换为下划线
    safe_name = re.sub(illegal_chars, '_', filename)
    # 去除首尾空格和点
    safe_name = safe_name.strip('. ')
    # 如果文件名为空，返回默认名称
    return safe_name or 'unnamed'


def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_safe_path_component(component):
    """
    安全地处理路径组件名称，保留中文字符
    """
    if not component:
        return 'unnamed'
    # 使用改进的文件名清理函数
    return sanitize_filename(component)


def get_user_display_name(user):
    """获取用户显示名称"""
    if hasattr(user, 'username'):
        return user.username
    elif hasattr(user, 'name'):
        return user.name
    return f"User_{user.id}" if hasattr(user, 'id') else "Unknown User"


# def create_upload_path(user_id, project_id, stage_id, task_id=None):
#     """创建并返回上传路径"""
#     try:
#         # 获取必要信息
#         user = User.query.get_or_404(user_id)
#         project = Project.query.get_or_404(project_id)
#         stage = ProjectStage.query.get_or_404(stage_id)
#
#         # 安全地处理各个路径组件，保留中文
#         safe_user = get_safe_path_component(user.username)
#         safe_project = get_safe_path_component(project.name)
#         safe_stage = get_safe_path_component(stage.name)
#
#         # 如果有任务ID，获取任务名称
#         safe_task = ''
#         if task_id:
#             task = StageTask.query.get_or_404(task_id)
#             safe_task = get_safe_path_component(task.name)
#
#         # 构建完整路径
#         if task_id:
#             relative_path = os.path.join(UPLOAD_FOLDER, safe_user, safe_project, safe_stage, safe_task)
#         else:
#             relative_path = os.path.join(UPLOAD_FOLDER, safe_user, safe_project, safe_stage)
#
#         absolute_path = os.path.join(current_app.root_path, relative_path)
#
#         # 确保目录存在
#         os.makedirs(absolute_path, exist_ok=True)
#
#         return relative_path, absolute_path
#     except Exception as e:
#         raise Exception(f"创建上传路径失败: {str(e)}")

# 2025年3月17日14:53:16
def create_upload_path(user_id, project_id, subproject_id, stage_id, task_id=None):
    """创建并返回上传路径，包含子项目层级"""
    try:
        # 获取必要信息
        user = User.query.get_or_404(user_id)
        project = Project.query.get_or_404(project_id)
        subproject = Subproject.query.get_or_404(subproject_id)
        stage = ProjectStage.query.get_or_404(stage_id)

        # 安全地处理各个路径组件，保留中文
        safe_user = get_safe_path_component(user.username)
        safe_project = get_safe_path_component(project.name)
        safe_subproject = get_safe_path_component(subproject.name)
        safe_stage = get_safe_path_component(stage.name)

        # 如果有任务ID，获取任务名称
        safe_task = ''
        if task_id:
            task = StageTask.query.get_or_404(task_id)
            safe_task = get_safe_path_component(task.name)

        # 构建完整路径，添加子项目层级
        if task_id:
            relative_path = os.path.join(UPLOAD_FOLDER, safe_user, safe_project, safe_subproject, safe_stage, safe_task)
        else:
            relative_path = os.path.join(UPLOAD_FOLDER, safe_user, safe_project, safe_subproject, safe_stage)

        absolute_path = os.path.join(current_app.root_path, relative_path)

        # 确保目录存在
        os.makedirs(absolute_path, exist_ok=True)

        return relative_path, absolute_path
    except Exception as e:
        raise Exception(f"创建上传路径失败: {str(e)}")


def generate_unique_filename(directory, original_filename):
    """生成唯一的文件名，处理文件名冲突"""
    base_name, extension = os.path.splitext(original_filename)
    safe_base_name = get_safe_path_component(base_name)
    counter = 1
    new_filename = f"{safe_base_name}{extension}"

    while os.path.exists(os.path.join(directory, new_filename)):
        new_filename = f"{safe_base_name}({counter}){extension}"
        counter += 1

    return new_filename


def get_user_display_name(user):
    """获取用户显示名称"""
    if hasattr(user, 'username'):
        return user.username
    elif hasattr(user, 'name'):
        return user.name
    return f"User_{user.id}" if hasattr(user, 'id') else "Unknown User"


def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_upload_path(user_name, project_name, stage_name):
    """生成上传路径"""
    # 移除文件名中的非法字符
    safe_user = secure_filename(user_name)
    safe_project = secure_filename(project_name)
    safe_stage = secure_filename(stage_name)

    # 构建路径
    path = os.path.join(BASE_UPLOAD_FOLDER, safe_user, safe_project, safe_stage)

    # 确保目录存在
    if not os.path.exists(path):
        os.makedirs(path)

    return path


def check_stage_progress(stage_id):
    """检查阶段进度是否达到100%"""
    stage = ProjectStage.query.get(stage_id)
    if not stage:
        return False
    return stage.progress >= 100


# 获取文件列表
# @files_bp.route('/stage/<int:stage_id>/task/<int:task_id>', methods=['GET'])
# @track_activity
# def get_stage_task_files(stage_id, task_id):
#     try:
#         files = ProjectFile.query.filter_by(
#             stage_id=stage_id,
#             task_id=task_id
#         ).all()
#
#         files_data = []
#         for file in files:
#             # 获取上传用户信息
#             uploader = User.query.get(file.upload_user_id)
#             uploader_name = get_user_display_name(uploader)
#
#             # 获取文件大小
#             try:
#                 file_size = os.path.getsize(os.path.join(current_app.root_path, file.file_path)) if os.path.exists(
#                     os.path.join(current_app.root_path, file.file_path)) else 0
#             except:
#                 file_size = 0
#
#             files_data.append({
#                 'id': file.id,
#                 'fileName': file.file_name,
#                 'originalName': file.original_name,
#                 'fileSize': file_size,
#                 'fileType': file.file_type,
#                 'uploadTime': file.upload_date.isoformat(),
#                 'uploader': uploader_name,
#                 'stageName': file.stage.name if file.stage else None,
#                 'taskName': file.task.name if file.task else None,
#                 'is_public': file.is_public
#             })
#
#         return jsonify(files_data)
#
#     except Exception as e:
#         print(f"获取阶段文件时出错： {str(e)}")
#         return jsonify({'error': str(e)}), 500


# 2025年3月17日14:54:14
@files_bp.route('/stage/<int:stage_id>/task/<int:task_id>', methods=['GET'])
@track_activity
def get_stage_task_files(stage_id, task_id):
    try:
        files = ProjectFile.query.filter_by(
            stage_id=stage_id,
            task_id=task_id
        ).all()

        files_data = []
        for file in files:
            # 获取上传用户信息
            uploader = User.query.get(file.upload_user_id)
            uploader_name = get_user_display_name(uploader)

            # 获取文件大小
            try:
                file_size = os.path.getsize(os.path.join(current_app.root_path, file.file_path)) if os.path.exists(
                    os.path.join(current_app.root_path, file.file_path)) else 0
            except:
                file_size = 0

            files_data.append({
                'id': file.id,
                'fileName': file.file_name,
                'originalName': file.original_name,
                'fileSize': file_size,
                'fileType': file.file_type,
                'uploadTime': file.upload_date.isoformat(),
                'uploader': uploader_name,
                'projectName': file.project.name if file.project else None,
                'subprojectName': file.subproject.name if file.subproject else None,  # 添加子项目名称
                'stageName': file.stage.name if file.stage else None,
                'taskName': file.task.name if file.task else None,
                'is_public': file.is_public
            })

        return jsonify(files_data)

    except Exception as e:
        print(f"获取阶段文件时出错： {str(e)}")
        return jsonify({'error': str(e)}), 500


# 获取子项目下所有文件 - 新增端点
@files_bp.route('/subproject/<int:subproject_id>/files', methods=['GET'])
@track_activity
def get_subproject_files(subproject_id):
    try:
        files = ProjectFile.query.filter_by(
            subproject_id=subproject_id
        ).all()

        files_data = []
        for file in files:
            # 获取上传用户信息
            uploader = User.query.get(file.upload_user_id)
            uploader_name = get_user_display_name(uploader)

            # 获取文件大小
            try:
                file_size = os.path.getsize(os.path.join(current_app.root_path, file.file_path)) if os.path.exists(
                    os.path.join(current_app.root_path, file.file_path)) else 0
            except:
                file_size = 0

            files_data.append({
                'id': file.id,
                'fileName': file.file_name,
                'originalName': file.original_name,
                'fileSize': file_size,
                'fileType': file.file_type,
                'uploadTime': file.upload_date.isoformat(),
                'uploader': uploader_name,
                'upload_user_id': file.upload_user_id,
                'projectName': file.project.name if hasattr(file, 'project') and file.project else None,
                'subprojectName': file.subproject.name if hasattr(file, 'subproject') and file.subproject else None,
                'stageName': file.stage.name if hasattr(file, 'stage') and file.stage else None,
                'taskName': file.task.name if hasattr(file, 'task') and file.task else None,
                'is_public': file.is_public
            })

        return jsonify(files_data)

    except Exception as e:
        print(f"获取子项目文件时出错： {str(e)}")
        return jsonify({'error': str(e)}), 500


# 获取阶段下所有文件 - 新增端点
@files_bp.route('/stage/<int:stage_id>/files', methods=['GET'])
@track_activity
def get_stage_files(stage_id):
    try:
        files = ProjectFile.query.filter_by(
            stage_id=stage_id
        ).all()

        files_data = []
        for file in files:
            # 获取上传用户信息
            uploader = User.query.get(file.upload_user_id)
            uploader_name = get_user_display_name(uploader)

            # 获取文件大小
            try:
                file_size = os.path.getsize(os.path.join(current_app.root_path, file.file_path)) if os.path.exists(
                    os.path.join(current_app.root_path, file.file_path)) else 0
            except:
                file_size = 0

            files_data.append({
                'id': file.id,
                'fileName': file.file_name,
                'originalName': file.original_name,
                'fileSize': file_size,
                'fileType': file.file_type,
                'uploadTime': file.upload_date.isoformat(),
                'uploader': uploader_name,
                'upload_user_id': file.upload_user_id,
                'projectName': file.project.name,
                'project_id': file.project_id,
                'project_name': file.project.name if hasattr(file, 'project') and file.project else None,
                'task_id': file.task_id,
                'subproject_id': file.subproject_id,
                'subprojectName': file.subproject.name,
                'stageName': file.stage.name,
                'taskName': file.task.name,
                'is_public': file.is_public,
            })

        return jsonify(files_data)

    except Exception as e:
        print(f"获取阶段文件时出错： {str(e)}")
        return jsonify({'error': str(e)}), 500


# 上传文件，带索引
# @files_bp.route('/<int:project_id>/stages/<int:stage_id>/tasks/<int:task_id>/upload', methods=['POST'])
# @track_activity
# def upload_task_file(project_id, stage_id, task_id):
#     file = request.files.get('file')
#     is_public = request.form.get('is_public', 'false').lower() == 'true'  # 获取是否公开的参数
#
#     if not file:
#         return jsonify({'error': '请提供文件'}), 400
#
#     # 验证任务、阶段和项目的关系
#     task = StageTask.query.get_or_404(task_id)
#     stage = ProjectStage.query.get_or_404(stage_id)
#
#     if task.stage_id != stage_id or stage.project_id != project_id:
#         return jsonify({'error': '任务、阶段或项目信息不匹配'}), 400
#
#     # 验证文件类型和大小
#     if not allowed_file(file.filename):
#         return jsonify({'error': '文件类型不允许'}), 400
#     if file.content_length > MAX_FILE_SIZE:
#         return jsonify({'error': '文件大小超过限制'}), 400
#
#     # 构建上传路径
#     employee_id = get_employee_id()
#     try:
#         relative_path, absolute_path = create_upload_path(employee_id, project_id, stage_id, task_id)
#     except Exception as e:
#         return jsonify({'error': f'创建上传路径失败: {str(e)}'}), 500
#
#     # 保存文件
#     try:
#         unique_filename = generate_unique_filename(absolute_path, file.filename)
#         file_path = os.path.join(absolute_path, unique_filename)
#         file.save(file_path)
#         mime_type = get_mime_type(file.filename) or file.content_type
#     except Exception as e:
#         return jsonify({'error': f'保存文件失败: {str(e)}'}), 500
#
#     # 创建文件记录
#     try:
#         project_file = ProjectFile(
#             project_id=project_id,
#             stage_id=stage_id,
#             task_id=task_id,
#             original_name=file.filename,
#             file_name=unique_filename,
#             file_type=mime_type,
#             file_path=os.path.join(relative_path, unique_filename),
#             upload_user_id=employee_id,
#             upload_date=datetime.now(),
#             text_extracted=False,
#             is_public=is_public  # 设置是否公开
#         )
#         db.session.add(project_file)
#         db.session.commit()
#
#         # 创建文件索引
#         try:
#             file_path = os.path.join(current_app.root_path, project_file.file_path)
#             if update_file_index(project_file.id, file_path, mime_type):
#                 print(f"已成功为 file 创建文件索引 {project_file.id}")
#             else:
#                 print(f"无法为 file 创建文件索引 {project_file.id}")
#         except Exception as e:
#             print(f"创建文件索引时出错: {str(e)}")
#
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({'error': f'数据库操作失败: {str(e)}'}), 500
#
#     return jsonify({'message': '文件上传成功', 'file_id': project_file.id})

# 2025年3月17日14:52:24
@files_bp.route('/<int:project_id>/subprojects/<int:subproject_id>/stages/<int:stage_id>/tasks/<int:task_id>/upload',
                methods=['POST'])
@track_activity
def upload_task_file(project_id, subproject_id, stage_id, task_id):
    file = request.files.get('file')
    is_public = request.form.get('is_public', 'false').lower() == 'true'  # 获取是否公开的参数

    if not file:
        return jsonify({'error': '请提供文件'}), 400

    # 验证任务、阶段、子项目和项目的关系
    task = StageTask.query.get_or_404(task_id)
    stage = ProjectStage.query.get_or_404(stage_id)
    subproject = Subproject.query.get_or_404(subproject_id)

    # 验证关系链
    if task.stage_id != stage_id or stage.subproject_id != subproject_id or subproject.project_id != project_id:
        return jsonify({'error': '任务、阶段、子项目或项目信息不匹配'}), 400

    # 验证文件类型和大小
    if not allowed_file(file.filename):
        return jsonify({'error': '文件类型不允许'}), 400
    if file.content_length > MAX_FILE_SIZE:
        return jsonify({'error': '文件大小超过限制'}), 400

    # 构建上传路径
    employee_id = get_employee_id()
    try:
        # 需要更新 create_upload_path 函数以支持子项目
        relative_path, absolute_path = create_upload_path(employee_id, project_id, subproject_id, stage_id, task_id)
    except Exception as e:
        return jsonify({'error': f'创建上传路径失败: {str(e)}'}), 500

    # 保存文件
    try:
        unique_filename = generate_unique_filename(absolute_path, file.filename)
        file_path = os.path.join(absolute_path, unique_filename)
        file.save(file_path)
        mime_type = get_mime_type(file.filename) or file.content_type
    except Exception as e:
        return jsonify({'error': f'保存文件失败: {str(e)}'}), 500

    # 创建文件记录
    try:
        project_file = ProjectFile(
            project_id=project_id,
            subproject_id=subproject_id,  # 添加子项目ID
            stage_id=stage_id,
            task_id=task_id,
            original_name=file.filename,
            file_name=unique_filename,
            file_type=mime_type,
            file_path=os.path.join(relative_path, unique_filename),
            upload_user_id=employee_id,
            upload_date=datetime.now(),
            text_extracted=False,
            is_public=is_public  # 设置是否公开
        )
        db.session.add(project_file)
        db.session.commit()

        # 创建文件索引
        try:
            file_path = os.path.join(current_app.root_path, project_file.file_path)
            if update_file_index(project_file.id, file_path, mime_type):
                print(f"已成功为 file 创建文件索引 {project_file.id}")
            else:
                print(f"无法为 file 创建文件索引 {project_file.id}")
        except Exception as e:
            print(f"创建文件索引时出错: {str(e)}")

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'数据库操作失败: {str(e)}'}), 500

    return jsonify({'message': '文件上传成功', 'file_id': project_file.id})


# 搜索功能，加权限展示，加公开属性
# 2025年3月17日14:55:30
@files_bp.route('/search', methods=['GET'])
@track_activity
def search_files():
    try:
        # 获取当前用户 ID 和角色
        employee_id = get_employee_id()
        current_user = User.query.get(employee_id)
        if not current_user:
            return jsonify({'error': '未找到用户'}), 404

        search_query = request.args.get('query', '').strip()
        visibility = request.args.get('visibility', '')  # 获取可见性筛选参数
        subproject_id = request.args.get('subproject_id', type=int)  # 获取子项目筛选参数

        if not search_query:
            return jsonify({'error': '搜索条件必填'}), 400

        base_query = ProjectFile.query.options(
            joinedload(ProjectFile.project),
            joinedload(ProjectFile.subproject),  # 添加子项目关联
            joinedload(ProjectFile.stage),
            joinedload(ProjectFile.task),
            joinedload(ProjectFile.upload_user),
            joinedload(ProjectFile.content)
        )

        # 子项目筛选
        if subproject_id:
            base_query = base_query.filter(ProjectFile.subproject_id == subproject_id)

        # 权限和可见性筛选逻辑保持不变
        if current_user.role != 1:  # 不是管理员
            if visibility == 'public':
                base_query = base_query.filter(ProjectFile.is_public == True)
            elif visibility == 'private':
                base_query = base_query.filter(
                    ProjectFile.upload_user_id == employee_id
                )
            else:
                base_query = base_query.filter(
                    or_(
                        ProjectFile.upload_user_id == employee_id,
                        ProjectFile.is_public == True
                    )
                )
        else:  # 管理员可以看到所有文件，但仍然应用可见性筛选
            if visibility == 'public':
                base_query = base_query.filter(ProjectFile.is_public == True)
            elif visibility == 'private':
                base_query = base_query.filter(ProjectFile.is_public == False)

        # 添加子项目到搜索条件
        search_conditions = [
            ProjectFile.original_name.ilike(f'%{search_query}%'),
            ProjectFile.file_name.ilike(f'%{search_query}%'),
            ProjectFile.file_type.ilike(f'%{search_query}%'),
            ProjectFile.file_path.ilike(f'%{search_query}%'),
            Project.name.ilike(f'%{search_query}%'),
            Subproject.name.ilike(f'%{search_query}%'),  # 添加子项目名称搜索
            ProjectStage.name.ilike(f'%{search_query}%'),
            StageTask.name.ilike(f'%{search_query}%'),
            User.username.ilike(f'%{search_query}%'),
            FileContent.content.ilike(f'%{search_query}%')
        ]

        search_results = base_query \
            .join(Project) \
            .join(Subproject) \
            .join(ProjectStage) \
            .join(StageTask) \
            .join(User, ProjectFile.upload_user_id == User.id) \
            .outerjoin(FileContent) \
            .filter(or_(*search_conditions)) \
            .all()

        results = []
        for file in search_results:
            try:
                file_path = os.path.join(current_app.root_path, file.file_path)
                file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

                result = {
                    'id': file.id,
                    'fileName': file.file_name,
                    'originalName': highlight_text(file.original_name, search_query),
                    'fileType': file.file_type,
                    'fileSize': file_size,
                    'uploadTime': file.upload_date.isoformat(),
                    'uploader': highlight_text(file.upload_user.username, search_query),
                    'projectName': highlight_text(file.project.name if file.project else None, search_query),
                    'subprojectName': highlight_text(file.subproject.name if file.subproject else None, search_query),
                    # 添加子项目名称
                    'stageName': highlight_text(file.stage.name if file.stage else None, search_query),
                    'taskName': highlight_text(file.task.name if file.task else None, search_query),
                    'is_public': file.is_public
                }

                # 添加内容预览
                if file.content:
                    preview = get_content_preview(
                        file.content.content,
                        search_query,
                        context_length=150
                    )
                    result['contentPreview'] = preview if preview else "无匹配内容"
                else:
                    result['contentPreview'] = "未提取内容"

                results.append(result)
            except Exception as e:
                print(f"处理文件 {file.id} 时出错: {str(e)}")
                continue

        return jsonify({
            'results': results,
            'total': len(results)
        })

    except Exception as e:
        print(f"搜索错误: {str(e)}")
        return jsonify({'error': str(e)}), 500


# 获取内容预览
def get_content_preview(content, query, context_length=150):
    """
    获取匹配内容的上下文预览，突出显示匹配的文本
    """
    if not content or not query:
        return None

    try:
        # 转换为小写进行不区分大小写的搜索
        content_lower = content.lower()
        query_lower = query.lower()

        # 查找匹配位置
        index = content_lower.find(query_lower)
        if index == -1:
            return None

        # 计算预览窗口的起始和结束位置
        start = max(0, index - context_length // 2)
        end = min(len(content), index + len(query) + context_length // 2)

        # 调整起始位置到单词边界（如果可能）
        while start > 0 and content[start - 1].isalnum():
            start -= 1

        # 调整结束位置到单词边界（如果可能）
        while end < len(content) - 1 and content[end].isalnum():
            end += 1

        # 构建预览文本
        preview = content[start:end].strip()

        # 添加省略号标记
        if start > 0:
            preview = f"...{preview}"
        if end < len(content):
            preview = f"{preview}..."

        return preview

    except Exception as e:
        print(f"生成内容预览时出错: {str(e)}")
        return None


def highlight_text(text, query):
    """
    在文本中为搜索关键词添加高亮标记
    使用特殊标记 {{highlight}} 和 {{/highlight}} 包裹匹配文本
    """
    if not text or not query:
        return text

    try:
        # 不区分大小写
        query_lower = query.lower()
        # 分割文本以保留原始大小写
        parts = []
        last_idx = 0
        text_lower = text.lower()

        while True:
            idx = text_lower.find(query_lower, last_idx)
            if idx == -1:
                parts.append(text[last_idx:])
                break

            parts.append(text[last_idx:idx])
            parts.append("{{highlight}}" + text[idx:idx + len(query)] + "{{/highlight}}")
            last_idx = idx + len(query)

        return "".join(parts)
    except Exception as e:
        print(f"高亮处理错误: {str(e)}")
        return text


#  文件删除接口
@files_bp.route('/<int:file_id>', methods=['DELETE'])
@track_activity
def delete_file(file_id):
    try:
        # 获取当前用户ID
        current_user_id = get_employee_id()

        # 获取文件记录
        file = ProjectFile.query.get_or_404(file_id)

        # 验证权限 (只允许文件上传者或管理员删除)
        user = User.query.get(current_user_id)
        if not user:
            return jsonify({'error': '用户未找到'}), 404

        if file.upload_user_id != current_user_id and user.role != 1:  # 1表示管理员角色
            return jsonify({'error': '没有权限删除此文件'}), 403

        # 获取物理文件路径
        file_path = os.path.join(current_app.root_path, file.file_path)

        # 删除物理文件
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except OSError as e:
                print(f"删除文件时出错： {e}")
                return jsonify({'error': '删除物理文件失败'}), 500

        # 检查并删除空文件夹
        try:
            directory = os.path.dirname(file_path)
            if os.path.exists(directory) and not os.listdir(directory):
                os.rmdir(directory)
        except OSError as e:
            print(f"删除空目录时出错： {e}")
            # 继续执行，因为这不是致命错误
        # 删除数据库记录
        db.session.delete(file)
        db.session.commit()

        return jsonify({
            'message': '文件删除成功',
            'file_id': file_id
        })

    except Exception as e:
        db.session.rollback()
        print(f"Error in delete_file: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ---------------------------------------------------------------------------------------------------------


# 文件下载
@files_bp.route('/download/<int:file_id>', methods=['GET'])
@track_activity
def download_file(file_id):
    try:
        file = ProjectFile.query.get_or_404(file_id)
        file_path = os.path.join(current_app.root_path, file.file_path)

        if not os.path.exists(file_path):
            return jsonify({'error': '文件不存在'}), 404

        return send_file(file_path, as_attachment=True, download_name=file.original_name)

    except Exception as e:
        print(f"下载错误：{str(e)}")
        return jsonify({'error': str(e)}), 500


# 从授权令牌中获取用户信息，并进行改进的错误处理
def get_user_info_from_token():
    try:
        auth_header = request.headers.get('Authorization', '')
        if not auth_header:
            print("无 Authorization 标头")
            return None

        token = auth_header.replace('Bearer ', '').strip()
        if not token:
            print("Empty token")
            return None

        payload = jwt.decode(
            token,
            current_app.config['SECRET_KEY'],
            algorithms=['HS256']
        )
        return payload

    except jwt.ExpiredSignatureError:
        print("令牌已过期")
        return None
    except jwt.InvalidTokenError as e:
        print(f"无效令牌错误： {str(e)}")
        return None
    except Exception as e:
        print(f"令牌处理中出现意外错误: {str(e)}")
        return None


#  获取带高亮标记的内容预览
def get_content_preview(content, query, context_length=150):
    if not content or not query:
        return None

    try:
        # 转换为小写进行不区分大小写的搜索
        content_lower = content.lower()
        query_lower = query.lower()

        # 查找匹配位置
        index = content_lower.find(query_lower)
        if index == -1:
            return None

        # 计算预览窗口的起始和结束位置
        start = max(0, index - context_length // 2)
        end = min(len(content), index + len(query) + context_length // 2)

        # 调整到单词边界
        while start > 0 and content[start - 1].isalnum():
            start -= 1
        while end < len(content) - 1 and content[end].isalnum():
            end += 1

        # 提取预览文本
        preview = content[start:end].strip()

        # 添加省略号
        if start > 0:
            preview = "..." + preview
        if end < len(content):
            preview = preview + "..."

        # 为预览文本添加高亮
        return highlight_text(preview, query)

    except Exception as e:
        print(f"生成预览错误: {str(e)}")
        return None


# 导出文件列表
@files_bp.route('/export', methods=['GET'])
@track_activity
def export_file_list():
    try:
        # 获取当前登录用户
        current_user_id = get_employee_id()
        current_user = User.query.get(current_user_id)
        if not current_user:
            return jsonify({'error': '未找到用户'}), 404

        # 构建基础查询
        base_query = db.session.query(
            ProjectFile,
            Project.name.label('project_name'),
            Subproject.name.label('subproject_name'),  # 添加子项目名称
            ProjectStage.name.label('stage_name'),
            StageTask.name.label('task_name'),
            User.username.label('uploader_name')
        ).join(
            Project, ProjectFile.project_id == Project.id
        ).join(
            Subproject, ProjectFile.subproject_id == Subproject.id  # 添加子项目关联
        ).join(
            ProjectStage, ProjectFile.stage_id == ProjectStage.id
        ).join(
            StageTask, ProjectFile.task_id == StageTask.id
        ).join(
            User, ProjectFile.upload_user_id == User.id
        )

        # 根据用户角色筛选数据
        if current_user.role != 1:  # 如果不是管理员，只能看到自己的文件
            base_query = base_query.filter(ProjectFile.upload_user_id == current_user_id)

        # 执行查询，并按项目、子项目、阶段、任务、上传时间排序
        files_query = base_query.order_by(
            Project.name,
            Subproject.name,  # 按子项目排序
            ProjectStage.name,
            StageTask.name,
            ProjectFile.upload_date
        ).all()

        # 统计信息查询（同样需要加入权限过滤）
        stats_query = db.session.query(
            func.count(ProjectFile.id).label('total_files')
        )
        if current_user.role != 1:
            stats_query = stats_query.filter(ProjectFile.upload_user_id == current_user_id)
        stats = stats_query.first()

        # 计算实际文件大小总和
        total_size = 0
        for file_data in files_query:
            file = file_data[0]  # ProjectFile对象在第一个位置
            file_path = os.path.join(current_app.root_path, file.file_path)
            if os.path.exists(file_path):
                total_size += os.path.getsize(file_path)

        # 创建Word文档
        doc = Document()

        # 设置中文字体
        style = doc.styles['Normal']
        doc.styles["Normal"].font.name = u"微软雅黑"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 设置标题字体
        for i in range(1, 5):  # 增加到5级标题，为子项目增加一级
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = '方正小标宋_GBK'
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')

        # 添加标题并设置其字体
        title = doc.add_heading(f'{current_user.username}的文件管理报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.name = '方正小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')

        doc.add_paragraph(f'导出时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'导出用户：{current_user.username}')
        doc.add_paragraph(f'文件总数：{stats.total_files or 0}个')
        doc.add_paragraph(f'文件总大小：{format_file_size(total_size)}')
        doc.add_paragraph()

        # 按项目、子项目、阶段、任务组织文件列表
        current_project = None
        current_subproject = None  # 添加子项目级别
        current_stage = None
        current_task = None
        current_table = None

        for file_data in files_query:
            file, project_name, subproject_name, stage_name, task_name, uploader = file_data

            # 项目层级
            if current_project != project_name:
                current_project = project_name
                heading = doc.add_heading(f'项目：{project_name}', level=1)
                # 确保每个标题的字体设置都正确
                for run in heading.runs:
                    run.font.name = '方正小标宋_GBK'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
                current_subproject = None
                current_stage = None
                current_task = None

            # 子项目层级
            if current_subproject != subproject_name:
                current_subproject = subproject_name
                heading = doc.add_heading(f'子项目：{subproject_name}', level=2)
                # 确保每个标题的字体设置都正确
                for run in heading.runs:
                    run.font.name = '方正小标宋_GBK'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
                current_stage = None
                current_task = None

            # 阶段层级
            if current_stage != stage_name:
                current_stage = stage_name
                heading = doc.add_heading(f'阶段：{stage_name}', level=3)
                # 确保每个标题的字体设置都正确
                for run in heading.runs:
                    run.font.name = '方正小标宋_GBK'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
                current_task = None

            # 任务层级
            if current_task != task_name:
                current_task = task_name
                heading = doc.add_heading(f'任务：{task_name}', level=4)
                # 确保每个标题的字体设置都正确
                for run in heading.runs:
                    run.font.name = '方正小标宋_GBK'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')

                # 添加文件列表表格
                current_table = doc.add_table(rows=1, cols=4)
                current_table.style = 'Table Grid'

                # 设置表格标题行字体
                header_cells = current_table.rows[0].cells
                for cell in header_cells:
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

                header_cells[0].text = '文件名'
                header_cells[1].text = '文件大小'
                header_cells[2].text = '上传时间'
                header_cells[3].text = '上传者'

            # 添加文件信息到表格
            file_path = os.path.join(current_app.root_path, file.file_path)
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            row_cells = current_table.add_row().cells

            # 设置表格内容字体
            for i, text in enumerate([
                file.original_name,
                format_file_size(file_size),
                file.upload_date.strftime("%Y-%m-%d %H:%M:%S"),
                uploader
            ]):
                paragraph = row_cells[i].paragraphs[0]
                run = paragraph.add_run(text)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # 保存文档到内存
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        # 生成安全的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'{current_user.username}_文件管理报告_{timestamp}.docx'
        filename = filename.encode('utf-8').decode('utf-8')  # 确保正确的UTF-8编码

        return send_file(
            doc_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"导出文件列表失败: {str(e)}")
        return jsonify({'error': str(e)}), 500


def format_file_size(size):
    """格式化文件大小"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            return f"{size:.2f} {unit}"
        size /= 1024
    return f"{size:.2f} TB"


# 2025年1月8日11:22:47
# 添加修改文件可见性的接口
@files_bp.route('/<int:file_id>/visibility', methods=['PUT'])
@track_activity
def update_file_visibility(file_id):
    try:
        # 获取当前用户
        current_user_id = get_employee_id()
        current_user = User.query.get(current_user_id)
        if not current_user:
            return jsonify({'error': '未找到用户'}), 404

        # 获取文件
        file = ProjectFile.query.get_or_404(file_id)

        # 验证权限(只有文件上传者或管理员可以修改)
        if file.upload_user_id != current_user_id and current_user.role != 1:
            return jsonify({'error': '没有权限修改此文件'}), 403

        # 获取请求数据
        data = request.get_json()
        is_public = data.get('is_public')

        if is_public is None:
            return jsonify({'error': '缺少必要参数'}), 400

        # 更新文件可见性
        file.is_public = is_public
        db.session.commit()

        # 记录操作日志
        activity_detail = f"{'设置文件为公开' if is_public else '设置文件为私有'}: {file.original_name}"
        UserActivityLog.log_activity(
            user_id=current_user_id,
            action_type='update_file_visibility',
            action_detail=activity_detail,
            resource_type='file',
            resource_id=file_id
        )

        return jsonify({
            'message': '文件可见性更新成功',
            'is_public': file.is_public
        })

    except Exception as e:
        db.session.rollback()
        print(f"更新文件可见性失败: {str(e)}")
        return jsonify({'error': str(e)}), 500


# 筛选所有公共文件
@files_bp.route('/public-files', methods=['GET'])
@track_activity
def get_public_files():
    try:
        # 获取当前用户 ID 和角色
        employee_id = get_employee_id()
        current_user = User.query.get(employee_id)
        if not current_user:
            return jsonify({'error': '未找到用户'}), 404

        # 查询所有公共文件
        public_files = ProjectFile.query.filter_by(is_public=True).all()

        files_data = []
        for file in public_files:
            # 获取上传用户信息
            uploader = User.query.get(file.upload_user_id)
            uploader_name = get_user_display_name(uploader)

            # 获取文件大小
            try:
                file_size = os.path.getsize(os.path.join(current_app.root_path, file.file_path)) if os.path.exists(
                    os.path.join(current_app.root_path, file.file_path)) else 0
            except:
                file_size = 0

            files_data.append({
                'id': file.id,
                'fileName': file.file_name,
                'originalName': file.original_name,
                'fileSize': file_size,
                'fileType': file.file_type,
                'uploadTime': file.upload_date.isoformat(),
                'uploader': uploader_name,
                'projectName': file.project.name if file.project else None,
                'subprojectName': file.subproject.name if file.subproject else None,  # 添加子项目名称
                'stageName': file.stage.name if file.stage else None,
                'taskName': file.task.name if file.task else None,
                'is_public': file.is_public
            })

        return jsonify(files_data)

    except Exception as e:
        print(f"获取公共文件时出错： {str(e)}")
        return jsonify({'error': str(e)}), 500


# 2025年1月15日17:11:19
# 已解决！！

def setup_fonts():
    """设置字体"""
    try:
        # 获取字体文件路径
        font_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts', 'simsun.ttf')

        # 注册基本字体
        pdfmetrics.registerFont(TTFont('SimSun', font_path))

        # 注册字体变体
        pdfmetrics.registerFontFamily(
            'SimSun',
            normal='SimSun',
            bold='SimSun',  # 使用同一个字体文件
            italic='SimSun',  # 使用同一个字体文件
            boldItalic='SimSun'  # 使用同一个字体文件
        )
        return True
    except Exception as e:
        print(f"字体设置失败: {str(e)}")
        return False


def clean_text(text):
    """清理文本内容"""
    if not text:
        return ""

    # 移除XML/HTML标签
    text = re.sub(r'<[^>]+>', '', text)

    # 转义HTML实体
    text = html.unescape(text)

    # 替换不可打印字符
    text = ''.join(char for char in text if char.isprintable())

    return text.strip()


def create_pdf_style():
    """创建PDF样式"""
    styles = getSampleStyleSheet()

    # 确保使用已注册的字体
    default_font = 'SimSun'

    # 基础样式
    basic_style = ParagraphStyle(
        'BasicStyle',
        parent=styles['Normal'],
        fontName=default_font,
        fontSize=12,
        leading=14,
        allowWidows=0,
        allowOrphans=0
    )

    # 标题样式
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=basic_style,
        fontName=default_font,
        fontSize=16,
        leading=20,
        spaceAfter=30,
        alignment=1
    )

    # 居中样式
    center_style = ParagraphStyle(
        'CenterStyle',
        parent=basic_style,
        fontName=default_font,
        alignment=1
    )

    # 右对齐样式
    right_style = ParagraphStyle(
        'RightStyle',
        parent=basic_style,
        fontName=default_font,
        alignment=2
    )

    return {
        'basic': basic_style,
        'title': title_style,
        'center': center_style,
        'right': right_style
    }


def extract_prefix_number(filename):
    """从文件名中提取前缀数字"""
    match = re.match(r'^(\d+)', filename)
    if match:
        return int(match.group(1))
    return float('inf')


def sort_files_by_prefix(files):
    """按照文件名前缀数字排序文件"""
    return sorted(files, key=lambda x: extract_prefix_number(x.file_name))


def should_use_landscape(data, font_name='SimSun', font_size=8):
    """
    判断是否应该使用横向布局
    通过计算内容宽度和竖向A4纸的可用宽度比较来决定
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    if not data or not data[0]:
        return False

    # 计算所有列的最大内容宽度
    max_col_widths = []
    for row in data:
        while len(max_col_widths) < len(row):
            max_col_widths.append(0)
        for i, cell in enumerate(row):
            content_width = stringWidth(str(cell), font_name, font_size)
            max_col_widths[i] = max(max_col_widths[i], content_width)

    # 计算总需要的宽度（包括一些边距和padding）
    total_required_width = sum(max_col_widths) + (len(max_col_widths) * 4)  # 4点的padding

    # A4纸的宽度（portrait模式）减去左右边距
    A4_PORTRAIT_WIDTH = A4[0] - 40  # 减去左右各20的边距

    # 如果需要的宽度超过竖向A4可用宽度的80%，建议使用横向
    return total_required_width > (A4_PORTRAIT_WIDTH * 0.8)


# Flask路由部分保持不变
merge_progress = {}
