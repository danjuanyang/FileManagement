# file_indexer.py
import os
import io

# from pymupdf import fitz

import openpyxl

# import sqlite3

# 使用条件导入
try:
    import fitz
except ImportError:
    try:
        from PyMuPDF import fitz
    except ImportError:
        fitz = None
# 添加一个简单的空调用
if False:
    fitz.Document()  # 永远不会执行，但 PyInstaller 能感知到

# 条件导入 chardet
try:
    import chardet
except ImportError:
    chardet = None

# 添加一个简单的空调用
if False:
    chardet.detect(b"")  # 永远不会执行，但 PyInstaller 能感知到


# 条件导入 docx
try:
    from docx import Document
except ImportError:
    Document = None

# docx 的空调用
if False:
    Document()  # 永远不会执行，但 PyInstaller 能感知到




from models import db, ProjectFile, FileContent


# def detect_file_encoding(file_path):
#     """检测文件编码"""
#     with open(file_path, 'rb') as file:
#         raw_data = file.read()
#         result = chardet.detect(raw_data)
#         return result['encoding']

def detect_file_encoding(file_path):
    """检测文件编码"""
    if chardet is None:
        print("编码检测模块未正确加载")
        return 'utf-8'  # 返回一个默认编码

    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        return result['encoding']


# def extract_text_from_docx(file_path):
#     """提取Word文档内容"""
#     try:
#         doc = Document(file_path)
#         full_text = []
#         for paragraph in doc.paragraphs:
#             full_text.append(paragraph.text)
#         return '\n'.join(full_text)
#     except Exception as e:
#         print(f"Extract docx error: {str(e)}")
#         return None

def extract_text_from_docx(file_path):
    """提取Word文档内容"""
    if Document is None:
        print("Word文档处理模块未正确加载")
        return None

    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Extract docx error: {str(e)}")
        return None




# def extract_text_from_pdf(file_path):
#     """提取PDF文档内容"""
#     try:
#         # 使用 PyMuPDF 替换 PyPDF2
#         doc = fitz.open(file_path)
#         text = []
#         for page in doc:
#             text.append(page.get_text())
#         doc.close()  # 记得关闭文档
#         return '\n'.join(text)
#     except Exception as e:
#         print(f"Extract PDF error: {str(e)}")
#         return None

# 2024年12月4日16:47:44
# 修改 extract_text_from_pdf 函数来处理 fitz 导入失败的情况
def extract_text_from_pdf(file_path):
    """提取PDF文档内容"""
    if fitz is None:
        print("PDF处理模块未正确加载")
        return None

    try:
        doc = fitz.open(file_path)
        text = []
        for page in doc:
            text.append(page.get_text())
        doc.close()
        return '\n'.join(text)
    except Exception as e:
        print(f"提取 PDF 错误： {str(e)}")
        return None


def extract_text_from_excel(file_path):
    """提取Excel文档内容"""
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.rows:
                row_text = ' '.join(str(cell.value) for cell in row if cell.value is not None)
                if row_text.strip():
                    text.append(row_text)
        return '\n'.join(text)
    except Exception as e:
        print(f"提取 Excel 错误： {str(e)}")
        return None


def extract_text_from_txt(file_path):
    """提取文本文件内容"""
    try:
        encoding = detect_file_encoding(file_path)
        with open(file_path, 'r', encoding=encoding) as file:
            return file.read()
    except Exception as e:
        print(f"提取 txt 错误： {str(e)}")
        return None


def create_file_index(file_path, file_type):
    """创建文件索引"""
    extractors = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': extract_text_from_docx,
        'application/pdf': extract_text_from_pdf,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': extract_text_from_excel,
        'text/plain': extract_text_from_txt
    }

    # 根据文件类型选择提取器
    extractor = extractors.get(file_type)
    if not extractor:
        print(f"不支持的文件类型： {file_type}")
        return None

    try:
        # 提取文件内容
        extracted_text = extractor(file_path)
        if extracted_text:
            # 处理提取的文本，移除多余的空白字符
            processed_text = ' '.join(extracted_text.split())
            return processed_text
        return None
    except Exception as e:
        print(f"索引创建错误：{str(e)}")
        return None


def update_file_index(project_file_id, file_path, file_type):
    """更新文件索引"""
    try:
        # 提取文件内容
        extracted_text = create_file_index(file_path, file_type)

        if extracted_text:
            # 获取或创建FileContent记录
            file_content = FileContent.query.filter_by(file_id=project_file_id).first()
            if not file_content:
                file_content = FileContent(file_id=project_file_id)

            file_content.content = extracted_text

            # 更新ProjectFile的text_extracted标志
            project_file = ProjectFile.query.get(project_file_id)
            project_file.text_extracted = True

            # 保存更改
            db.session.add(file_content)
            db.session.add(project_file)
            db.session.commit()

            return True
    except Exception as e:
        print(f"索引更新错误： {str(e)}")
        db.session.rollback()
        return False


# 文件类型映射
MIME_TYPE_MAPPING = {
    'doc': 'application/msword',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'pdf': 'application/pdf',
    'xls': 'application/vnd.ms-excel',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'txt': 'text/plain'
}


def get_mime_type(filename):
    """根据文件扩展名获取MIME类型"""
    ext = filename.rsplit('.', 1)[-1].lower()
    return MIME_TYPE_MAPPING.get(ext)
